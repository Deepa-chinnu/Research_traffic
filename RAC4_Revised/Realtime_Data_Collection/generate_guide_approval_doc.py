"""
Generate a Word document for guide approval - Real-Time Data Collection Roadmap
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ============================================================================
# PAGE SETUP
# ============================================================================
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ============================================================================
# STYLES
# ============================================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ============================================================================
# TITLE PAGE
# ============================================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('REAL-TIME TRAFFIC DATA COLLECTION\nROADMAP & APPROVAL DOCUMENT')
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('For PhD Research')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(127, 140, 141)

doc.add_paragraph()

research_title = doc.add_paragraph()
research_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = research_title.add_run('"Predict and Mitigate Road Traffic Flow\nUsing Machine Learning"')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(192, 57, 43)

doc.add_paragraph()
doc.add_paragraph()

# Scholar info
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('Scholar: Deepa C\nReg. No: 2390019')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(44, 62, 80)

doc.add_paragraph()
doc.add_paragraph()

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run('March 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(127, 140, 141)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_heading('TABLE OF CONTENTS', level=1)
doc.add_paragraph()

toc_items = [
    ('1.', 'Purpose of This Document'),
    ('2.', 'Current Research Status (What Has Been Done)'),
    ('3.', 'Why Real-Time Data Is Needed'),
    ('4.', 'Proposed Data Collection Approach'),
    ('5.', 'API Selection & Justification'),
    ('6.', 'Data Parameters Mapping'),
    ('7.', 'Collection Schedule & Duration'),
    ('8.', 'Cost Analysis'),
    ('9.', 'Technical Implementation Plan'),
    ('10.', 'Data Preprocessing Pipeline'),
    ('11.', 'How Real-Time Data Integrates With Existing Research'),
    ('12.', 'Risk Assessment & Mitigation'),
    ('13.', 'Timeline'),
    ('14.', 'Expected Outcomes'),
    ('15.', 'Approval Checklist'),
]

for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.size = Pt(11)

doc.add_page_break()

# ============================================================================
# SECTION 1: PURPOSE
# ============================================================================
doc.add_heading('1. Purpose of This Document', level=1)

doc.add_paragraph(
    'This document presents the complete roadmap for collecting real-time traffic data '
    'from Bangalore roads using publicly available APIs. The purpose is to seek guide '
    'approval before starting the data collection process, which will run for a minimum '
    'of 3 months to gather sufficient data for the final PhD thesis.'
)

doc.add_paragraph(
    'The document covers: what APIs will be used, what data will be collected, how it '
    'maps to the existing Kaggle dataset parameters, the cost (zero), the technical '
    'implementation, and the timeline for completion.'
)

# ============================================================================
# SECTION 2: CURRENT STATUS
# ============================================================================
doc.add_heading('2. Current Research Status', level=1)

doc.add_heading('2.1 Work Completed So Far', level=2)

completed_items = [
    ('Literature Review', '25 research papers reviewed across neural network and ML-based traffic prediction approaches (Bartlett et al., Li et al., Zhou et al., Zou et al., Mihaita et al.)'),
    ('Dataset Analysis', 'Bangalore Traffic Dataset from Kaggle (8,936 records, 16 columns, 8 areas, 16 roads, 2022-2024) fully analyzed with EDA and temporal analysis'),
    ('Individual ML Models', '10 ML model families trained with 27+ hyperparameter configurations: Linear Regression, Ridge, Lasso, ElasticNet, SVR, KNN, Decision Tree, Random Forest, Extra Trees, Gradient Boosting'),
    ('Model Comparison', 'Comprehensive comparison using R\u00b2, RMSE, MAE, MAPE metrics with utility score ranking'),
    ('Model Improvements', 'Hyperparameter tuning and optimization for each model'),
    ('Novel ODTC Framework', 'Developed the Online Dynamic Temporal Context (ODTC) ML Framework - a 5-stage pipeline (MSTFE + DCE + AFS + MES + OAM) achieving R\u00b2 = 1.000000'),
    ('Documentation', '220+ output files including plots, reports, flow diagrams, and complete research documentation'),
]

for title, desc in completed_items:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('2.2 What Is Missing for Final Thesis', level=2)

doc.add_paragraph(
    'The current research uses a Kaggle dataset which is sufficient for methodology '
    'demonstration and framework development. However, for the final PhD thesis, '
    'real-time data from actual Bangalore roads is required to:'
)

missing_items = [
    'Validate the ODTC Framework on real-world data (not just Kaggle)',
    'Demonstrate practical applicability of the research',
    'Show the framework works with live, continuously changing traffic patterns',
    'Strengthen the thesis with verifiable, reproducible real-time data',
    'Meet the PhD committee requirement for real-world data validation',
]

for item in missing_items:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================================
# SECTION 3: WHY REAL-TIME DATA
# ============================================================================
doc.add_heading('3. Why Real-Time Data Is Needed', level=1)

doc.add_heading('3.1 Limitations of Kaggle Dataset', level=2)

limitations = [
    ('Origin Unknown', 'The Kaggle dataset does not provide details about how the data was collected (sensors, surveys, or simulated).'),
    ('Data Quality Concern', 'The Environmental Impact column has a perfect correlation (r = +1.0) with Traffic Volume, causing data leakage. This column and 4 other fabricated columns have been removed from the pipeline.'),
    ('Not Verifiable', 'We cannot verify the accuracy of the data against ground truth since the collection methodology is unknown.'),
    ('Static Dataset', 'The dataset is fixed (2022-2024) and cannot demonstrate the ODTC Framework\'s online adaptation capability with genuinely new data.'),
]

for title, desc in limitations:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('3.2 Benefits of Real-Time Data', level=2)

benefits = [
    ('Verifiable Source', 'Data comes from established APIs (TomTom, HERE) used by millions of navigation users worldwide. The data source is transparent and reproducible.'),
    ('Live Traffic Patterns', 'Captures actual current traffic conditions on the same 16 Bangalore roads from the original study.'),
    ('Online Adaptation Testing', 'The ODTC Framework\'s Stage 5 (Online Adaptation Module) can be genuinely tested with incrementally arriving real data.'),
    ('Research Credibility', 'Using recognized traffic data providers (TomTom, HERE) adds credibility to the thesis.'),
    ('Reproducibility', 'Other researchers can replicate the data collection using the same APIs and locations.'),
]

for title, desc in benefits:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================================
# SECTION 4: PROPOSED APPROACH
# ============================================================================
doc.add_heading('4. Proposed Data Collection Approach', level=1)

doc.add_paragraph(
    'We propose using three complementary traffic and weather APIs to collect data '
    'for the same 16 Bangalore road locations used in the Kaggle dataset. The APIs '
    'are queried every 15 minutes, and the raw data is aggregated to daily records '
    'matching the exact format of the original dataset.'
)

doc.add_heading('4.1 High-Level Architecture', level=2)

doc.add_paragraph(
    'The data collection system follows this pipeline:'
)

pipeline_steps = [
    'Step 1: Automated Python script runs every 15 minutes (via Windows Task Scheduler)',
    'Step 2: For each of the 16 roads, query TomTom (speed, travel time), HERE (congestion), and OpenWeatherMap (weather)',
    'Step 3: Raw data is appended to a CSV file with timestamps',
    'Step 4: Periodically, a preprocessing script aggregates 15-minute data to daily records',
    'Step 5: Daily records have the same 16 columns as the Kaggle dataset',
    'Step 6: The processed daily data can be directly used with the existing ODTC Framework',
]

for i, step in enumerate(pipeline_steps):
    doc.add_paragraph(step, style='List Number')

doc.add_heading('4.2 Key Design Decisions', level=2)

decisions = [
    ('Same 16 Roads', 'We collect data for the exact same 16 road/intersection locations as the Kaggle dataset to enable direct comparison between Kaggle results and real-time results.'),
    ('Same 8 Areas', 'Indiranagar, M.G. Road, Koramangala, Whitefield, Jayanagar, Hebbal, Yeshwanthpur, Electronic City.'),
    ('15-Minute Intervals', 'This captures intra-day traffic patterns (peak hours, off-peak, night) and provides 96 data points per road per day for robust daily aggregation.'),
    ('Multiple APIs', 'Using both TomTom and HERE provides redundancy and cross-validation of speed/congestion measurements.'),
    ('Daily Aggregation', 'The final output matches the daily granularity of the Kaggle dataset, ensuring compatibility with all existing ML models.'),
]

for title, desc in decisions:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================================
# SECTION 5: API SELECTION
# ============================================================================
doc.add_heading('5. API Selection & Justification', level=1)

# API comparison table
table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['API', 'Data Provided', 'Free Tier Limit', 'Our Daily Usage', 'Cost']
for i, header in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

api_data = [
    ['TomTom\nTraffic Flow', 'Current speed (km/h)\nFree-flow speed\nTravel time\nRoad closures', '2,500\nrequests/day', '1,536\nrequests/day', '$0'],
    ['HERE\nTraffic Flow', 'Speed (km/h)\nFree-flow speed\njamFactor (0-10)\nTraversability', '250,000\nrequests/month', '46,080\nrequests/month', '$0'],
    ['OpenWeatherMap', 'Weather condition\nTemperature\nHumidity, Wind\nRain, Visibility', '1,000\ncalls/day', '96\ncalls/day', '$0'],
    ['TomTom\nIncidents', 'Incident type\nDelay & severity\nConstruction info', 'Included in\n2,500/day', '~200\ncalls/day', '$0'],
]

for i, row_data in enumerate(api_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

doc.add_paragraph()

doc.add_heading('5.1 Why These APIs?', level=2)

doc.add_paragraph(
    'TomTom: Ranks Bengaluru as the 2nd most congested city globally in their Traffic Index. '
    'Excellent coverage of Bangalore roads. Used by millions of users for navigation, ensuring '
    'reliable probe data. No credit card required for the free tier.'
)

doc.add_paragraph(
    'HERE: Has 96% map coverage in India with a 3,000-person local team. Provides the unique '
    'jamFactor metric (0.0 = free flow to 10.0 = fully blocked) which directly maps to our '
    'Congestion Level parameter. No credit card required.'
)

doc.add_paragraph(
    'OpenWeatherMap: The most widely used weather API globally. Provides weather condition '
    'categories that map directly to our dataset\'s Weather Conditions column (Clear, Rain, '
    'Fog, Windy). Daily limit of 1,000 calls far exceeds our 96 calls/day need.'
)

doc.add_heading('5.2 APIs Considered But Not Selected', level=2)

not_selected = [
    ('Google Maps Routes API', 'Provides traffic-aware travel time but does NOT return actual speed values in km/h (only categorical labels: NORMAL, SLOW, TRAFFIC_JAM). Also costs $10/1,000 requests with traffic data, which would be ~$370 for 3 months.'),
    ('Ola Maps API', 'India-specific and free, but only provides ETA/duration with traffic, not raw speed or congestion metrics. No jamFactor equivalent.'),
    ('Mappls (MapmyIndia)', 'Good India coverage but paid plans are $300-600/month. Not suitable for a 3-month free collection.'),
]

for title, desc in not_selected:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================================
# SECTION 6: DATA PARAMETERS MAPPING
# ============================================================================
doc.add_heading('6. Data Parameters Mapping', level=1)

doc.add_paragraph(
    'The following table shows how each of the 16 columns in the Kaggle dataset '
    'will be obtained or derived from the real-time APIs:'
)

# Parameter mapping table
table = doc.add_table(rows=20, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

param_headers = ['Column Name', 'Source', 'Method', 'Type']
for i, header in enumerate(param_headers):
    cell = table.rows[0].cells[i]
    cell.text = header
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

param_data = [
    ['Date', 'System Clock', 'Collection timestamp (IST)', 'Direct'],
    ['Area Name', 'Config File', 'Predefined road-to-area mapping', 'Direct'],
    ['Road/Intersection Name', 'Config File', '16 GPS coordinate pairs', 'Direct'],
    ['Traffic Volume', 'TomTom + HERE', 'Estimated using Greenshields traffic flow model from speed + congestion', 'Derived'],
    ['Average Speed', 'TomTom', 'currentSpeed field (km/h) - direct measurement', 'Direct'],
    ['Travel Time Index', 'TomTom', 'currentTravelTime / freeFlowTravelTime ratio', 'Direct'],
    ['Congestion Level', 'HERE', 'jamFactor (0-10) scaled to 0-100', 'Direct'],
    ['Road Capacity Utilization', 'TomTom', '(1 - currentSpeed/freeFlowSpeed) x 100', 'Derived'],
    ['Incident Reports', 'TomTom Incidents', 'Count of incidents within 500m radius', 'Direct'],
    ['Weather Conditions', 'OpenWeatherMap', 'Mapped to: Clear/Rain/Fog/Windy/Overcast', 'Direct'],
    ['Roadwork & Construction', 'TomTom Incidents', 'Filtered for construction-type incidents', 'Direct'],
    ['Temperature', 'OpenWeatherMap', 'Daily mean temperature (Celsius)', 'Direct'],
    ['Humidity', 'OpenWeatherMap', 'Daily mean humidity (%)', 'Direct'],
    ['Wind Speed', 'OpenWeatherMap', 'Daily mean wind speed (m/s)', 'Direct'],
    ['Rain Volume', 'OpenWeatherMap', 'Daily mean rainfall (mm/h)', 'Direct'],
    ['Visibility', 'OpenWeatherMap', 'Daily mean visibility (m)', 'Direct'],
    ['Free Flow Speed', 'TomTom', 'Free-flow speed (km/h)', 'Direct'],
    ['Jam Factor', 'TomTom', 'Congestion factor 0-10', 'Direct'],
    ['Confidence', 'TomTom', 'Speed measurement confidence', 'Direct'],
]

for i, row_data in enumerate(param_data):
    for j, cell_text in enumerate(row_data):
        cell = table.rows[i+1].cells[j]
        cell.text = cell_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(8)

doc.add_paragraph()

doc.add_heading('6.1 Important Note on Traffic Volume', level=2)

doc.add_paragraph(
    'No publicly available API provides actual vehicle counts (Traffic Volume). '
    'All commercial traffic APIs (Google, TomTom, HERE, Ola) derive their data from '
    'GPS probe data (anonymized traces from smartphones and connected vehicles). They '
    'estimate speed and congestion but do not count individual vehicles.'
)

doc.add_paragraph(
    'For actual vehicle counts, physical sensors (inductive loop detectors, cameras with '
    'computer vision) would be needed, which is beyond the scope of this research.'
)

doc.add_paragraph(
    'We estimate Traffic Volume using the Greenshields traffic flow model:'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Volume = Density x Speed, where Density = k_jam x (1 - Speed/FreeFlowSpeed)')
run.bold = True
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(192, 57, 43)

doc.add_paragraph(
    'This is a well-established traffic engineering model used in transportation research. '
    'The estimated volumes will be validated against the Kaggle dataset ranges to ensure '
    'reasonableness.'
)

doc.add_heading('6.2 Columns Removed from Pipeline', level=2)

doc.add_paragraph(
    'Five columns from the original Kaggle dataset have been removed from the real-time '
    'data collection pipeline because they were fabricated with random noise and had no '
    'real API source:'
)

removed_cols = [
    'Environmental Impact (had a perfect r=+1.0 correlation with Traffic Volume, causing data leakage)',
    'Public Transport Usage (fabricated with random noise)',
    'Traffic Signal Compliance (fabricated with random noise)',
    'Parking Usage (fabricated with random noise)',
    'Pedestrian and Cyclist Count (fabricated with random noise)',
]

for item in removed_cols:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'These have been replaced with real API-sourced columns: Temperature, Humidity, '
    'Wind Speed, Rain Volume, Visibility, Free Flow Speed, Jam Factor, and Confidence.'
)

# ============================================================================
# SECTION 7: COLLECTION SCHEDULE
# ============================================================================
doc.add_heading('7. Collection Schedule & Duration', level=1)

schedule_table = doc.add_table(rows=8, cols=2)
schedule_table.style = 'Light Grid Accent 1'

schedule_data = [
    ['Parameter', 'Value'],
    ['Collection Frequency', 'Every 15 minutes (96 times per day)'],
    ['Roads Monitored', '16 roads across 8 Bangalore areas'],
    ['Raw Records Per Day', '16 roads x 96 intervals = 1,536 records'],
    ['Daily Records (Aggregated)', '16 records per day (1 per road)'],
    ['Minimum Duration', '3 months (90 days)'],
    ['Total Raw Records (3 months)', '138,240 records'],
    ['Total Daily Records (3 months)', '1,440 records (comparable to Kaggle\'s 8,936)'],
]

for i, (col1, col2) in enumerate(schedule_data):
    schedule_table.rows[i].cells[0].text = col1
    schedule_table.rows[i].cells[1].text = col2
    if i == 0:
        for cell in schedule_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('7.1 Why 15-Minute Intervals?', level=2)

doc.add_paragraph(
    '15-minute intervals capture the full daily traffic cycle: morning peak (8-10 AM), '
    'off-peak (11 AM - 4 PM), evening peak (5-8 PM), and night (9 PM - 7 AM). This '
    'provides rich intra-day data for robust daily aggregation and enables future '
    'research on hourly/sub-daily prediction if needed.'
)

# ============================================================================
# SECTION 8: COST
# ============================================================================
doc.add_heading('8. Cost Analysis', level=1)

p = doc.add_paragraph()
run = p.add_run('Total Cost for 3 Months of Data Collection: $0 (ZERO)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(39, 174, 96)

doc.add_paragraph()

cost_table = doc.add_table(rows=5, cols=4)
cost_table.style = 'Light Grid Accent 1'

cost_data = [
    ['API', 'Our Monthly Usage', 'Free Monthly Limit', 'Monthly Cost'],
    ['TomTom', '46,080 calls', '75,000 calls', '$0'],
    ['HERE', '46,080 calls', '250,000 calls', '$0'],
    ['OpenWeatherMap', '2,880 calls', '30,000 calls', '$0'],
    ['TOTAL', '', '', '$0'],
]

for i, row_data in enumerate(cost_data):
    for j, cell_text in enumerate(row_data):
        cost_table.rows[i].cells[j].text = cell_text
        if i == 0 or i == 4:
            for paragraph in cost_table.rows[i].cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()
doc.add_paragraph(
    'No credit card is required for TomTom and HERE. OpenWeatherMap\'s free tier '
    'also does not require payment information for the basic Current Weather API.'
)

# ============================================================================
# SECTION 9: TECHNICAL IMPLEMENTATION
# ============================================================================
doc.add_heading('9. Technical Implementation Plan', level=1)

doc.add_heading('9.1 Software Requirements', level=2)

requirements = [
    'Python 3.8+ (already installed)',
    'requests library (pip install requests)',
    'pandas and numpy (already installed for ML work)',
    'Windows Task Scheduler (built into Windows)',
    'Internet connection (for API calls)',
]

for item in requirements:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('9.2 Files Created', level=2)

files_table = doc.add_table(rows=6, cols=2)
files_table.style = 'Light Grid Accent 1'

files_data = [
    ['File', 'Purpose'],
    ['config.py', 'API keys, 16 road GPS coordinates, collection settings'],
    ['realtime_data_collector.py', 'Main script - calls TomTom + HERE + OpenWeatherMap for all 16 roads'],
    ['preprocess_realtime_data.py', 'Converts raw 15-min data to daily format matching Kaggle columns'],
    ['setup_scheduler.py', 'Sets up Windows Task Scheduler for automated collection'],
    ['API_SETUP_GUIDE.txt', 'Step-by-step instructions to get free API keys'],
]

for i, (col1, col2) in enumerate(files_data):
    files_table.rows[i].cells[0].text = col1
    files_table.rows[i].cells[1].text = col2
    if i == 0:
        for cell in files_table.rows[i].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('9.3 16 Road Locations', level=2)

loc_table = doc.add_table(rows=17, cols=4)
loc_table.style = 'Light Grid Accent 1'

loc_headers = ['Road/Intersection', 'Area', 'Latitude', 'Longitude']
for i, h in enumerate(loc_headers):
    loc_table.rows[0].cells[i].text = h
    for paragraph in loc_table.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

roads = [
    ('100 Feet Road', 'Indiranagar', '12.9719', '77.6412'),
    ('CMH Road', 'Indiranagar', '12.9686', '77.6403'),
    ('Trinity Circle', 'M.G. Road', '12.9716', '77.6197'),
    ('Anil Kumble Circle', 'M.G. Road', '12.9780', '77.5990'),
    ('Sony World Junction', 'Koramangala', '12.9340', '77.6165'),
    ('Sarjapur Road', 'Koramangala', '12.9100', '77.6680'),
    ('Marathahalli Bridge', 'Whitefield', '12.9591', '77.7010'),
    ('ITPL Main Road', 'Whitefield', '12.9857', '77.7263'),
    ('Jayanagar 4th Block', 'Jayanagar', '12.9250', '77.5830'),
    ('South End Circle', 'Jayanagar', '12.9390', '77.5770'),
    ('Hebbal Flyover', 'Hebbal', '13.0358', '77.5970'),
    ('Ballari Road', 'Hebbal', '13.0100', '77.5750'),
    ('Tumkur Road', 'Yeshwanthpur', '13.0220', '77.5510'),
    ('Yeshwanthpur Circle', 'Yeshwanthpur', '13.0270', '77.5540'),
    ('Hosur Road', 'Electronic City', '12.8920', '77.6460'),
    ('Silk Board Junction', 'Electronic City', '12.9173', '77.6229'),
]

for i, (road, area, lat, lon) in enumerate(roads):
    loc_table.rows[i+1].cells[0].text = road
    loc_table.rows[i+1].cells[1].text = area
    loc_table.rows[i+1].cells[2].text = lat
    loc_table.rows[i+1].cells[3].text = lon
    for j in range(4):
        for paragraph in loc_table.rows[i+1].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# ============================================================================
# SECTION 10: PREPROCESSING
# ============================================================================
doc.add_heading('10. Data Preprocessing Pipeline', level=1)

doc.add_paragraph(
    'The raw 15-minute interval data is preprocessed into daily records using '
    'the following pipeline:'
)

preprocess_steps = [
    'Load raw CSV with 15-minute interval records',
    'Group by (Date, Road Name) to create daily aggregates',
    'Average Speed: Mean of all TomTom and HERE speed readings for that day',
    'Travel Time Index: Mean ratio of current travel time to free-flow travel time',
    'Congestion Level: Mean of HERE jamFactor readings, scaled from 0-10 to 0-100',
    'Road Capacity Utilization: Derived from speed ratio (1 - currentSpeed/freeFlowSpeed) x 100',
    'Traffic Volume: Estimated using Greenshields model with daily speed and congestion',
    'Weather: Most frequent weather condition during the day',
    'Incidents & Roadwork: Maximum incident count, any construction flag',
    'Output: CSV with same 16 columns as Kaggle dataset',
]

for i, step in enumerate(preprocess_steps):
    doc.add_paragraph(step, style='List Number')

# ============================================================================
# SECTION 11: INTEGRATION
# ============================================================================
doc.add_heading('11. How Real-Time Data Integrates With Existing Research', level=1)

doc.add_paragraph(
    'The real-time data collection is designed to seamlessly integrate with all existing '
    'research work:'
)

integration = [
    ('Compatible Columns', 'The preprocessed daily CSV contains the core Kaggle columns plus real API-sourced features, compatible with all existing Python scripts (step1 through step7).'),
    ('Same Road Names', 'The 16 road/intersection names are identical to the Kaggle dataset.'),
    ('Same Area Names', 'The 8 area names are identical to the Kaggle dataset.'),
    ('ODTC Framework Validation', 'The real-time daily data can be fed directly into the ODTC Framework (step7_ODTC_ML_Framework.py) to validate its performance on real-world data.'),
    ('Online Adaptation Testing', 'Stage 5 of the ODTC Framework (Online Adaptation Module) can be genuinely tested with incrementally arriving real data, proving the framework\'s ability to adapt.'),
    ('Comparative Analysis', 'A direct comparison between Kaggle results and real-time results can be presented in the thesis, showing the framework\'s generalizability.'),
]

for title, desc in integration:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

# ============================================================================
# SECTION 12: RISK ASSESSMENT
# ============================================================================
doc.add_heading('12. Risk Assessment & Mitigation', level=1)

risk_table = doc.add_table(rows=7, cols=3)
risk_table.style = 'Light Grid Accent 1'

risk_data = [
    ['Risk', 'Impact', 'Mitigation'],
    ['API downtime', 'Missing data for some intervals', 'Using 2 traffic APIs (TomTom + HERE) provides redundancy. Missing intervals are handled gracefully in daily aggregation.'],
    ['Internet outage', 'No data collection during outage', 'Task Scheduler automatically resumes when internet returns. Daily aggregation handles gaps.'],
    ['API free tier changes', 'May need to pay or switch API', 'Current usage is well below free limits. Multiple APIs provide backup options.'],
    ['Computer sleep/shutdown', 'Collection pauses', 'Configure Windows power settings to prevent sleep. Use Task Scheduler with "Wake computer" option.'],
    ['GPS coordinate mismatch', 'Data for wrong road segment', 'Coordinates verified against Google Maps. TomTom and HERE snap to nearest road.'],
    ['Rate limiting', 'Some requests rejected', 'Built-in delays (0.3s between calls) and error retry logic in the collector script.'],
]

for i, row_data in enumerate(risk_data):
    for j, cell_text in enumerate(row_data):
        risk_table.rows[i].cells[j].text = cell_text
        for paragraph in risk_table.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

# ============================================================================
# SECTION 13: TIMELINE
# ============================================================================
doc.add_heading('13. Timeline', level=1)

timeline_table = doc.add_table(rows=8, cols=3)
timeline_table.style = 'Light Grid Accent 1'

timeline_data = [
    ['Phase', 'Duration', 'Activities'],
    ['Setup', 'Day 1', 'Register for API keys, configure scripts, test connections'],
    ['Pilot Collection', 'Days 2-7', 'Run collector for 1 week, verify data quality, fix any issues'],
    ['Full Collection', 'Months 1-3', 'Automated collection every 15 minutes for 3 months'],
    ['Weekly Preprocessing', 'Every Sunday', 'Run preprocessing script, check quality report'],
    ['Mid-point Review', 'Month 1.5', 'Review data quality, present interim results to guide'],
    ['Final Processing', 'After Month 3', 'Final preprocessing, generate complete daily dataset'],
    ['Validation', 'Week after', 'Run ODTC Framework on real-time data, compare with Kaggle results'],
]

for i, row_data in enumerate(timeline_data):
    for j, cell_text in enumerate(row_data):
        timeline_table.rows[i].cells[j].text = cell_text
        for paragraph in timeline_table.rows[i].cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
                if i == 0:
                    run.bold = True

# ============================================================================
# SECTION 14: EXPECTED OUTCOMES
# ============================================================================
doc.add_heading('14. Expected Outcomes', level=1)

outcomes = [
    '1,440+ daily traffic records for 16 Bangalore roads (3 months)',
    '138,240+ raw 15-minute interval records for detailed analysis',
    'Validated ODTC Framework performance on real-world data',
    'Comparison study: Kaggle data vs Real-time data results',
    'Demonstrated online adaptation capability with genuinely new data',
    'Reproducible data collection methodology documented for thesis',
    'Strengthened thesis with verifiable, real-time data validation',
]

for item in outcomes:
    doc.add_paragraph(item, style='List Bullet')

# ============================================================================
# SECTION 15: APPROVAL CHECKLIST
# ============================================================================
doc.add_heading('15. Approval Checklist', level=1)

doc.add_paragraph(
    'Please review and approve the following items before data collection begins:'
)

doc.add_paragraph()

checklist = [
    'Data Collection Approach: Using TomTom + HERE + OpenWeatherMap APIs',
    'Road Locations: Same 16 roads and 8 areas as Kaggle dataset',
    'Collection Duration: Minimum 3 months of automated collection',
    'Collection Frequency: Every 15 minutes (96 data points per road per day)',
    'Cost: $0 (all APIs within free tier)',
    'Traffic Volume Estimation: Using Greenshields model (since no API provides actual vehicle counts)',
    'Removed Columns: 5 fabricated columns dropped; replaced with real API-sourced weather and traffic features',
    'Data Format: Core columns from Kaggle dataset plus real API-sourced features',
    'Timeline: Setup in 1 day, pilot for 1 week, full collection for 3 months',
    'Integration: Real-time data will be validated using existing ODTC Framework',
]

for item in checklist:
    p = doc.add_paragraph()
    run = p.add_run('\u2610  ')  # Checkbox character
    run.font.size = Pt(14)
    p.add_run(item)

doc.add_paragraph()
doc.add_paragraph()

# Signature lines
doc.add_paragraph('_' * 40)
p = doc.add_paragraph()
run = p.add_run('Guide Signature & Date')
run.font.color.rgb = RGBColor(127, 140, 141)

doc.add_paragraph()
doc.add_paragraph('_' * 40)
p = doc.add_paragraph()
run = p.add_run('Scholar Signature & Date')
run.font.color.rgb = RGBColor(127, 140, 141)

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Remarks / Suggestions:')
run.bold = True
doc.add_paragraph()
doc.add_paragraph('_' * 60)
doc.add_paragraph()
doc.add_paragraph('_' * 60)
doc.add_paragraph()
doc.add_paragraph('_' * 60)

# ============================================================================
# SAVE
# ============================================================================
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Realtime_Data_Collection_Approval_Document.docx')
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"Pages: ~15 pages")
